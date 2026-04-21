"""Document text extraction for PDF, TXT, and DOCX files.

Each parser returns a list of ParsedPage with text and page number,
providing a unified interface for downstream chunking.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    page_number: int
    text: str


@dataclass
class ParsedDocument:
    pages: list[ParsedPage] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def page_count(self) -> int:
        return len(self.pages)


def parse_pdf(file_path: str | Path) -> ParsedDocument:
    import pymupdf

    doc = pymupdf.open(str(file_path))
    pages = []
    try:
        for i, page in enumerate(doc):
            text = page.get_text("text")
            pages.append(ParsedPage(page_number=i + 1, text=text))
    finally:
        doc.close()
    return ParsedDocument(pages=pages)


def parse_txt(file_path: str | Path) -> ParsedDocument:
    text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)])


def parse_docx(file_path: str | Path) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    return ParsedDocument(pages=[ParsedPage(page_number=1, text=full_text)])


def parse_html(file_path: str | Path) -> ParsedDocument:
    """Extract main article text from an HTML snapshot on disk.

    Uses trafilatura (boilerplate-removal + readability heuristics). Falls
    back to a BeautifulSoup body-text dump if trafilatura extracts nothing,
    which sometimes happens on very short or unusually structured pages.
    """
    raw = Path(file_path).read_bytes()

    import trafilatura

    text = trafilatura.extract(
        raw,
        include_comments=False,
        include_tables=True,
        favor_precision=False,
        no_fallback=False,
    )

    if not text or not text.strip():
        # Fallback: crude body text via BeautifulSoup.
        from bs4 import BeautifulSoup

        try:
            soup = BeautifulSoup(raw, "lxml")
        except Exception:
            soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript", "template"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)

    return ParsedDocument(pages=[ParsedPage(page_number=1, text=text or "")])


PARSERS = {
    "pdf": parse_pdf,
    "txt": parse_txt,
    "docx": parse_docx,
    "html": parse_html,
}


def parse_file(file_path: str | Path, file_type: str) -> ParsedDocument:
    """Parse a file based on its type, returning structured text with page info."""
    parser = PARSERS.get(file_type)
    if not parser:
        raise ValueError(f"No parser for file type: {file_type}")
    logger.info("Parsing %s file: %s", file_type, file_path)
    result = parser(file_path)
    logger.info("Parsed %d pages, %d chars", result.page_count, len(result.full_text))
    return result
